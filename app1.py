import streamlit as st
from openai import OpenAI
import PyPDF2 
import re
from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO
import os
import json

# --- 1. CONFIGURATIE ---
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

def laad_alle_cvs():
    cv_map = {}
    folder = "cv_database"
    for bestand in os.listdir(folder):
        if bestand.endswith(".pdf"):
            with open(os.path.join(folder, bestand), "rb") as f:
                reader = PyPDF2.PdfReader(f)
                tekst = "".join([p.extract_text() for p in reader.pages])
                naam = bestand.replace(".pdf", "").replace("_", " ")
                cv_map[naam] = tekst
    return cv_map

# Functie om de tekst goed in de template te krijgen (Behouden zoals gevraagd)
def create_formatted_docx(text, is_cv=True):
    doc = Document()
    lines = text.split('\n')
    
    for line in lines:
        clean_line = line.replace('*', '').replace('#', '').strip()
        if not clean_line: 
            continue
            
        if is_cv and clean_line.startswith('-'):
            p = doc.add_paragraph(style='List Bullet')
            run_text = clean_line.lstrip('-').strip()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_together = True 
        else:
            p = doc.add_paragraph()
            run_text = clean_line
            p.paragraph_format.space_after = Pt(6)

        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0

        run = p.add_run(run_text)
        run.font.name = 'Poppins Light'
        run.font.size = Pt(9)
        
        if is_cv:
            upper_line = clean_line.upper().strip()
            headers = ["KERNCOMPETENTIES", "WERKERVARING", "OPLEIDING", "CURSUSSEN & TRAININGEN", "VAARDIGHEDEN & COMPETENTIES", "RELEVANTE ERVARING"]
            if any(header in upper_line for header in headers) or \
               ("|" in clean_line and "InTheArena" in clean_line) or \
               (re.search(r'\(\d{4}\s*-\s*.*\)', clean_line)) or \
               (re.search(r'\(\d{4}\)', clean_line)):
                run.bold = True
        elif ":" in clean_line and len(clean_line) < 50:
            run.bold = True

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 2. LOGIN & PAGINA BEHEER ---
if 'authenticated' not in st.session_state:
    st.session_state.authenticated = False
if 'page' not in st.session_state:
    st.session_state.page = "home"

# State voor versiebeheer
if 'cv_versions' not in st.session_state:
    st.session_state.cv_versions = []
# State voor motivatie en analyse per versie
if 'mot_versions' not in st.session_state:
    st.session_state.mot_versions = []
if 'ana_versions' not in st.session_state:
    st.session_state.ana_versions = []

if 'preload_cv_tekst' not in st.session_state:
    st.session_state.preload_cv_tekst = None
if 'preload_opdracht' not in st.session_state:
    st.session_state.preload_opdracht = None
if 'preload_naam' not in st.session_state:
    st.session_state.preload_naam = None

if not st.session_state.authenticated:
    st.title("InTheArena Portaal")
    password = st.text_input("Voer het wachtwoord in:", type="password")
    if st.button("Log in"):
        if password == st.secrets["APP_PASSWORD"]:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("Onjuist wachtwoord.")
    st.stop()

# --- 3. HOMEPAGE ---
if st.session_state.page == "home":
    st.title("Welkom bij het InTheArena Portaal")
    st.write("Maak een keuze uit de onderstaande tools:")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("📄 CV & Motivatie Generator", use_container_width=True):
            st.session_state.page = "cv_builder"
            st.rerun()
            
    with col2:
        if st.button("🎯 Test geschiktheid opdracht", use_container_width=True):
            st.session_state.page = "geschiktheid_test"
            st.rerun()

# --- 4. PAGINA: CV BUILDER ---
elif st.session_state.page == "cv_builder":
    if st.sidebar.button("⬅ Terug naar Menu"):
        st.session_state.page = "home"
        st.rerun()

    st.title("InTheArena CV Builder")

    # --- Automatisch laden vanuit geschiktheidstest ---
    if st.session_state.preload_cv_tekst:
        st.info(f"📋 CV van **{st.session_state.preload_naam}** is voorgeladen vanuit de geschiktheidstest.")
        
        with st.spinner(f"CV van {st.session_state.preload_naam} wordt automatisch herschreven..."):
            cv_text = st.session_state.preload_cv_tekst
            job_description_auto = st.session_state.preload_opdracht
    
            cv_system = (
                "Jij bent mijn AI-assistent for het professionaliseren van CV's voor brokerportalen. "
                "Jouw taak is om een nieuw, volledig herschreven CV te genereren in exact dezelfde structuur, "
                "layout, tone-of-voice en schrijfstijl als het originele InTheArena-format.\n\n"
                "BELANGRIJK: De lengte van het herschreven CV MOET tussen de 700 en 900 woorden liggen. "
                "Wees specifiek in resultaten en verantwoordelijkheden.\n\n"
                "BELANGRIJK: Zorg dat de 'Harde Eisen' van de opdrachtomschrijving LETTERLIJK terugkomen in de tekst.\n"
                "INSTRUCTIES VOOR INHOUD:\n"
                "- Herschrijf slim, nooit verzinnen: Gebruik ALLEEN werk dat daadwerkelijk in het originele CV staat.\n"
                "- Kwaliteiten InTheArena: workshops faciliteren, analyse en structuur aanbrengen, "
                "communiceren en overtuigen, gedrag en teams begeleiden, implementatie realiseren, resultaten meten en borgen.\n"
                "- Schrijf 100% op basis van de uitvraag.\n\n"
                "GEWENSTE STRUCTUUR:\n"
                "Naam | Consultant | InTheArena en daaronder twee alinea's over de kracht en aanpak\n"
                "Kerncompetenties (met bullets)\n"
                "Relevante ervaring t.o.v. functie [Naam Functie] (met bullets)\n"
                "Werkervaring (Functie | Bedrijf (Jaartal), met daaronder bullets)\n"
                "Opleiding\n"
                "Cursussen & trainingen\n"
                "Vaardigheden en competenties\n"
                "GEBRUIK VOOR ELKE BULLET een streepje (-) gevolgd door een tab.\n"
                "GEBRUIK GEEN ASTERISKEN *"
            )
    
            cv_res = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "system", "content": cv_system},
                          {"role": "user", "content": f"Opdracht: {job_description_auto}\n\nCV: {cv_text}"}],
                temperature=0.3
            )
            st.session_state.cv_versions.append(cv_res.choices[0].message.content)
    
            mot_res = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "system", "content": "Schrijf een korte motivatie (200-300 woorden) in InTheArena-stijl. Geen asterisken (*)."},
                          {"role": "user", "content": f"Opdracht: {job_description_auto}\n\nCV: {cv_text}"}],
                temperature=0.4
            )
            st.session_state.mot_versions.append(mot_res.choices[0].message.content)
    
            ana_res = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "system", "content": "Analyseer ontbrekende vaardigheden kritisch. Geen asterisken (*)."},
                          {"role": "user", "content": f"Opdracht: {job_description_auto}\n\nCV: {cv_text}"}],
                temperature=0.2
            )
            st.session_state.ana_versions.append(ana_res.choices[0].message.content)
    
            # Preload leegmaken zodat het niet opnieuw triggert
            st.session_state.preload_cv_tekst = None
            st.session_state.preload_opdracht = None
            st.session_state.preload_naam = None
            st.rerun()
    
    # Normale upload flow (alleen tonen als er geen preload actief is)
    uploaded_file = st.file_uploader("Upload het originele CV (PDF)", type="pdf")
    job_description = st.text_area("Plak hier de opdracht van de klant:", height=200)

    if st.button("Genereer Documenten"):
        if uploaded_file and job_description:
            with st.spinner('Bezig met herschrijven...'):
                reader = PyPDF2.PdfReader(uploaded_file)
                cv_text = "".join([page.extract_text() for page in reader.pages])
                
                cv_system = (
                    "Jij bent mijn AI-assistent for het professionaliseren van CV's voor brokerportalen. "
                    "Jouw taak is om een nieuw, volledig herschreven CV te genereren in exact dezelfde structuur, "
                    "layout, tone-of-voice en schrijfstijl als het originele InTheArena-format.\n\n"
                    "BELANGRIJK: De lengte van het herschreven CV MOET tussen de 700 en 900 woorden liggen. "
                    "Breid de beschrijvingen van de werkervaring uit op basis van het origineel om dit te bereiken. "
                    "Wees specifiek in resultaten en verantwoordelijkheden.\n\n"
                    "BELANGRIJK: Zorg dat de 'Harde Eisen' of eisen waar de kandidaat aan moet voldoen van de opdrachtomschrijving LETTERLIJK terugkomen in de tekst, bij voorkeur in de 'Relevante ervaring' en onder de specifieke functies in de 'Werkervaring'.\n"
                    "INSTRUCTIES VOOR INHOUD:\n"
                    "- Herschrijf slim, nooit verzinnen: Gebruik ALLEEN werk dat daadwerkelijk in het originele CV staat.\n"
                    "- Kwaliteiten InTheArena: workshops faciliteren, analyse en structuur aanbrengen, "
                    "communiceren en overtuigen, gedrag en teams begeleiden, implementatie realiseren, resultaten meten en borgen.\n"
                    "- Schrijf 100% op basis van de uitvraag.\n\n"
                    "GEWENSTE STRUCTUUR:\n"
                    "Naam | Consultant | InTheArena en daaronder twee alinea's over de kracht en aanpak\n"
                    "Kerncompetenties (met bullets)\n"
                    "Relevante ervaring t.o.v. functie [Naam Functie] (met bullets)\n"
                    "Werkervaring (Functie | Bedrijf (Jaartal), met daaronder bullets)\n"
                    "Opleiding\n"
                    "Cursussen & trainingen\n"
                    "Vaardigheden en competenties\n"
                    "GEBRUIK VOOR ELKE BULLET een streepje (-) gevolgd door een tab.\n"
                    "GEBRUIK GEEN ASTERISKEN *"
                )
                
                cv_res = client.chat.completions.create(
                    model="gpt-4o", 
                    messages=[{"role": "system", "content": cv_system},
                              {"role": "user", "content": f"Opdracht: {job_description}\n\nCV: {cv_text}"}],
                    temperature=0.3
                )
                st.session_state.cv_versions.append(cv_res.choices[0].message.content)
    
                mot_res = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[{"role": "system", "content": "Schrijf een korte motivatie (200-300 woorden) in InTheArena-stijl. Geen asterisken (*)."},
                              {"role": "user", "content": f"Opdracht: {job_description}\n\nCV: {cv_text}"}],
                    temperature=0.4
                )
                st.session_state.mot_versions.append(mot_res.choices[0].message.content)
    
                ana_res = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[{"role": "system", "content": "Analyseer ontbrekende vaardigheden kritisch. Geen asterisken (*)."},
                              {"role": "user", "content": f"Opdracht: {job_description}\n\nCV: {cv_text}"}],
                    temperature=0.2
                )
                st.session_state.ana_versions.append(ana_res.choices[0].message.content)
                st.rerun()
    
    # Feedback & Versiebeheer sectie
    if st.session_state.cv_versions:
        st.divider()
        st.subheader("📚 Versiebeheer & Verfijnen")
        
        version_options = [f"Versie {i+1}" for i in range(len(st.session_state.cv_versions))]
        selected_version_index = st.selectbox("Selecteer een CV versie om te bekijken of te updaten:", 
                                            options=range(len(version_options)), 
                                            format_func=lambda x: version_options[x],
                                            index=len(version_options)-1)
        
        selected_cv = st.session_state.cv_versions[selected_version_index]
        selected_mot = st.session_state.mot_versions[selected_version_index] if selected_version_index < len(st.session_state.mot_versions) else None
        selected_ana = st.session_state.ana_versions[selected_version_index] if selected_version_index < len(st.session_state.ana_versions) else None

        feedback = st.text_area("Wat moet er veranderd worden aan de geselecteerde versie?", placeholder="Bijv: Meer focus op projectmanagement.")
        
        if st.button("Update CV"):
            if feedback:
                with st.spinner('CV, motivatiebrief en analyse worden bijgewerkt...'):
                    # Update CV
                    cv_update = client.chat.completions.create(
                        model="gpt-4o", 
                        messages=[
                            {"role": "system", "content": (
                                "Jij bent een redacteur. Je krijgt een volledig CV en feedback. "
                                "Jouw taak is om het VOLLEDIGE CV opnieuw uit te spugen (700-900 woorden), "
                                "waarbij je EXACT de structuur en de rest van de tekst behoudt, "
                                "en ENKEL de specifieke wijzigingen uit de feedback doorvoert. "
                                "Verwijder geen secties! Gebruik GEEN asterisken (*)."
                            )},
                            {"role": "user", "content": f"Huidig CV:\n{selected_cv}\n\nFeedback van de gebruiker: {feedback}"}
                        ],
                        temperature=0.3
                    )
                    new_cv = cv_update.choices[0].message.content
                    st.session_state.cv_versions.append(new_cv)

                    # Automatisch nieuwe motivatiebrief genereren op basis van het bijgewerkte CV
                    mot_update = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[
                            {"role": "system", "content": "Schrijf een korte motivatie (200-300 woorden) in InTheArena-stijl op basis van het bijgewerkte CV. Geen asterisken (*)."},
                            {"role": "user", "content": f"Bijgewerkt CV:\n{new_cv}"}
                        ],
                        temperature=0.4
                    )
                    st.session_state.mot_versions.append(mot_update.choices[0].message.content)

                    # Automatisch nieuwe analyse genereren op basis van het bijgewerkte CV
                    ana_update = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[
                            {"role": "system", "content": "Analyseer ontbrekende vaardigheden kritisch op basis van het bijgewerkte CV. Geen asterisken (*)."},
                            {"role": "user", "content": f"Bijgewerkt CV:\n{new_cv}"}
                        ],
                        temperature=0.2
                    )
                    st.session_state.ana_versions.append(ana_update.choices[0].message.content)

                    st.success("Nieuwe versie aangemaakt — inclusief bijgewerkte motivatiebrief en analyse!")
                    st.rerun()

        # Download sectie gebaseerd op de geselecteerde versie
        st.divider()
        col1, col2, col3 = st.columns(3)
        with col1:
            st.download_button(f"Download CV ({version_options[selected_version_index]})", 
                              data=create_formatted_docx(selected_cv, True), 
                              file_name=f"Herschreven_CV_V{selected_version_index+1}.docx")
        with col2:
            if selected_mot:
                st.download_button(f"Download Motivatie ({version_options[selected_version_index]})", 
                                   data=create_formatted_docx(selected_mot, False), 
                                   file_name=f"Motivatie_V{selected_version_index+1}.docx")
        with col3:
            if selected_ana:
                st.download_button(f"Download Analyse ({version_options[selected_version_index]})", 
                                   data=create_formatted_docx(selected_ana, False), 
                                   file_name=f"Analyse_V{selected_version_index+1}.docx")
        
        st.info("### Preview Analyse")
        if selected_ana:
            st.markdown(selected_ana.replace('*', ''))

# --- 5. PAGINA: GESCHIKTHEID TEST (Upload & Analyse) ---
elif st.session_state.page == "geschiktheid_test":
    if st.sidebar.button("⬅ Terug naar Menu"):
        st.session_state.page = "home"
        st.rerun()

    st.title("🎯 Test geschiktheid opdracht")
    cv_database = laad_alle_cvs()

    if not cv_database:
        st.error("⚠️ Geen CV's gevonden in de map `cv_database/`. Voeg PDF-bestanden toe in je GitHub repo.")
        st.stop()

    st.success(f"✅ {len(cv_database)} CV('s) geladen: {', '.join(cv_database.keys())}")
    st.divider()

    opdracht = st.text_area("Plak hier de opdracht / functieomschrijving:", height=250,
                            placeholder="Bijv: Wij zoeken een ervaren projectmanager met kennis van Agile/Scrum...")

    st.subheader("Welke kandidaten wil je toetsen?")
    alle_namen = list(cv_database.keys())
    geselecteerde_kandidaten = st.multiselect(
        "Selecteer kandidaten (laat leeg = iedereen):",
        options=alle_namen, default=alle_namen
    )

    if st.button("🔍 Analyseer geschiktheid", type="primary"):
        if not opdracht:
            st.warning("Vul eerst een opdracht in.")
        elif not geselecteerde_kandidaten:
            st.warning("Selecteer minimaal één kandidaat.")
        else:
            resultaten = []
            progress = st.progress(0, text="Bezig met analyseren...")
            totaal = len(geselecteerde_kandidaten)

            for i, naam in enumerate(geselecteerde_kandidaten):
                progress.progress(i / totaal, text=f"Analyseer {naam}...")
                systeem_prompt = (
                    "Jij bent een kritische HR-specialist bij InTheArena. "
                    "Beoordeel of de consultant geschikt is voor de opdracht, beoordeel op harde eisen als daar niet aan voldaan wordt voldoet de kandidaat niet. "
                    "Beredeneer: Verbind specifieke projecten of resultaten uit het CV aan de opdracht.\n"
                    "InTheArena-factor: Let op ervaring met workshops, implementatie en structuur.\n\n"
                    "Geef je analyse UITSLUITEND als JSON zonder extra tekst:\n"
                    '{"score": <0-100>, "advies": "<Geschikt / Mogelijk geschikt / Niet geschikt>", '
                    '"sterke_punten": ["...", "..."], "tekortkomingen": ["...", "..."], '
                    '"samenvatting": "<2-3 zinnen>"}\n'
                    "Wees eerlijk. Verzin niets wat niet in het CV staat."
                )
                try:
                    response = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[
                            {"role": "system", "content": systeem_prompt},
                            {"role": "user", "content": f"Opdracht:\n{opdracht}\n\nCV van {naam}:\n{cv_database[naam]}"}
                        ],
                        temperature=0.2
                    )
                    raw = response.choices[0].message.content.strip()
                    clean = re.sub(r"```json|```", "", raw).strip()
                    data = json.loads(clean)
                    data["naam"] = naam
                    resultaten.append(data)
                except Exception as e:
                    resultaten.append({"naam": naam, "score": 0, "advies": "Fout",
                                       "sterke_punten": [], "tekortkomingen": [],
                                       "samenvatting": f"Kon niet analyseren: {e}"})

            progress.progress(1.0, text="Klaar!")
            resultaten.sort(key=lambda x: x.get("score", 0), reverse=True)

            st.divider()
            st.subheader("📊 Resultaten")
            for r in resultaten:
                score = r.get("score", 0)
                kleur = "🟢" if score >= 70 else ("🟡" if score >= 45 else "🔴")
                with st.expander(f"{kleur} **{r['naam']}** — {score}/100 | {r.get('advies','')}", expanded=True):
                    st.markdown(f"**📝 Samenvatting:** {r.get('samenvatting', '-')}")
                    col1, col2 = st.columns(2)
                    with col1:
                        st.markdown("**✅ Sterke punten:**")
                        for p in r.get("sterke_punten", []): st.markdown(f"- {p}")
                    with col2:
                        st.markdown("**⚠️ Tekortkomingen:**")
                        for p in r.get("tekortkomingen", []): st.markdown(f"- {p}")
            
                    st.divider()
                    if st.button(f"✍️ Herschrijf CV van {r['naam']} voor deze opdracht", key=f"herschrijf_{r['naam']}"):
                        st.session_state.preload_cv_tekst = cv_database[r['naam']]
                        st.session_state.preload_opdracht = opdracht
                        st.session_state.preload_naam = r['naam']
                        # Reset versies zodat je met een schone lei begint
                        st.session_state.cv_versions = []
                        st.session_state.mot_versions = []
                        st.session_state.ana_versions = []
                        st.session_state.page = "cv_builder"
                        st.rerun()

